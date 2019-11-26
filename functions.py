import docx


# Function to parse through the text in a specified docx file Will add all bold elements to a list that will be returned
def get_text(filename):
    doc = docx.Document(filename)
    bold_text = []
    for para in doc.paragraphs:
        # If paragraph is
        for run in para.runs:
            if run.bold:
                bold_text.append(para.text)
            elif run.bold and run.underline:
                bold_text.append(f"{para.text}\n")
    return bold_text
    # return '\n'.join(bold_text)


# Will write a doc with all the bold words returned from the get_text function
def write_doc(text_list, file_name):
    doc = docx.Document()
    for words in text_list:
        doc.add_paragraph(words)
        doc.add_paragraph('\n')
    doc.save(f"{file_name}.docx")
